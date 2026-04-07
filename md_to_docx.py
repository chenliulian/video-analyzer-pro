#!/usr/bin/env python3
"""
将 Markdown 文件转换为 Word 文档
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 未安装 python-docx")
    print("请运行: pip install python-docx")
    sys.exit(1)


def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


import re

def parse_inline_formatting(text):
    """解析行内格式，返回格式片段列表"""
    # 处理粗体 **text**
    pattern = r'\*\*(.*?)\*\*'
    parts = []
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # 添加普通文本
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        # 添加粗体文本
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    # 添加剩余文本
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    return parts if parts else [('normal', text)]


def parse_markdown(md_content):
    """简单解析 Markdown 内容"""
    lines = md_content.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith('# '):
            result.append(('h1', line[2:]))
        elif line.startswith('## '):
            result.append(('h2', line[3:]))
        elif line.startswith('### '):
            result.append(('h3', line[4:]))
        elif line.startswith('#### '):
            result.append(('h4', line[5:]))
        # 引用
        elif line.startswith('> '):
            result.append(('quote', line[2:]))
        # 分隔线
        elif line.strip() == '---':
            result.append(('hr', ''))
        # 列表项
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            result.append(('bullet', line.strip()[2:]))
        # 表格行
        elif line.startswith('|') and '---' not in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '') == '' for c in cells):
                result.append(('table_row', cells))
        # 普通段落
        elif line.strip():
            result.append(('para', line))
        
        i += 1
    
    return result


def md_to_docx(md_file, output_file=None):
    """将 Markdown 文件转换为 Word 文档"""
    md_path = Path(md_file)
    if not md_path.exists():
        print(f"错误: 文件不存在 {md_file}")
        return False
    
    if output_file is None:
        output_file = md_path.with_suffix('.docx')
    
    # 读取 Markdown 内容
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 解析 Markdown
    elements = parse_markdown(md_content)
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    for elem_type, content in elements:
        if elem_type == 'h1':
            p = doc.add_heading(content, level=1)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 18, True)
        
        elif elem_type == 'h2':
            p = doc.add_heading(content, level=2)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 16, True)
        
        elif elem_type == 'h3':
            p = doc.add_heading(content, level=3)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 14, True)
        
        elif elem_type == 'h4':
            p = doc.add_paragraph()
            parts = parse_inline_formatting(content)
            for fmt, text in parts:
                run = p.add_run(text)
                set_chinese_font(run, 'SimHei', 12, True)
        
        elif elem_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            parts = parse_inline_formatting(content)
            for fmt, text in parts:
                run = p.add_run(text)
                set_chinese_font(run, 'SimSun', 11, bold=(fmt == 'bold'))
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        elif elem_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_inline_formatting(content)
            for fmt, text in parts:
                run = p.add_run(text)
                set_chinese_font(run, 'SimSun', 11, bold=(fmt == 'bold'))
        
        elif elem_type == 'para':
            p = doc.add_paragraph()
            parts = parse_inline_formatting(content)
            for fmt, text in parts:
                run = p.add_run(text)
                set_chinese_font(run, 'SimSun', 12, bold=(fmt == 'bold'))
        
        elif elem_type == 'hr':
            doc.add_paragraph('_' * 50)
    
    # 保存文档
    doc.save(output_file)
    print(f"✓ 转换完成: {output_file}")
    return True


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <markdown_file> [output_file]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    md_to_docx(md_file, output_file)
