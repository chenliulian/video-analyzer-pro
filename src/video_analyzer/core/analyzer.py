#!/usr/bin/env python3
"""
视频分析器 Pro 版 - 完整自动化处理流程
功能: 提取帧 → 提取音频 → 语音识别 → OCR识别 → LLM精炼 → 生成PDF/Word
"""

import os
import sys
import ssl
import json
import subprocess
from pathlib import Path
from typing import List, Tuple, Dict, Optional

try:
    from moviepy import VideoFileClip
    import whisper
    from PIL import Image
    import cv2
    import numpy as np
except ImportError as e:
    print(f"错误: 缺少必要的依赖包 - {e}")
    print("安装方法: pip3 install -r requirements.txt")
    sys.exit(1)

# 导入本地模块
from ..ocr.qwen_ocr import QwenOCR
from ..llm.refine_agent import TextRefineAgent
from ..asr.qwen_asr import QwenASR

# 解决SSL证书问题
ssl._create_default_https_context = ssl._create_unverified_context

# 尝试导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: 未安装 python-docx,无法生成Word文档")
    print("安装: pip3 install python-docx")


class VideoAnalyzerPro:
    """视频分析器Pro版 - 完整自动化流程"""
    
    def __init__(self, 
                 video_file: str,
                 use_llm: bool = True,
                 llm_config: Optional[Dict] = None,
                 ocr_config: Optional[Dict] = None,
                 whisper_model: str = "base",
                 asr_engine: str = "whisper",
                 asr_config: Optional[Dict] = None,
                 skip_existing: bool = True):
        """
        初始化视频分析器
        
        参数:
            video_file: 视频文件路径
            use_llm: 是否使用LLM精炼
            llm_config: LLM配置字典
            ocr_config: OCR配置字典
            whisper_model: Whisper模型大小
            asr_engine: 语音识别引擎 ("whisper" 或 "qwen")
            asr_config: ASR配置字典
            skip_existing: 是否跳过已存在的文件
        """
        # 检查视频文件
        if not os.path.exists(video_file):
            raise FileNotFoundError(f"视频文件不存在: {video_file}")
        
        self.video_file = video_file
        self.video_id = Path(video_file).stem
        self.use_llm = use_llm
        self.whisper_model_name = whisper_model
        self.asr_engine = asr_engine
        self.skip_existing = skip_existing
        
        # 目录设置
        self.results_dir = Path("results")
        self.frames_dir = self.results_dir / "extracted_frames"
        self.results_dir.mkdir(exist_ok=True)
        self.frames_dir.mkdir(exist_ok=True)
        
        # 文件路径
        self.audio_file = self.results_dir / f"{self.video_id}.mp3"
        self.transcript_file = self.results_dir / f"{self.video_id}_transcript.txt"
        self.ocr_file = self.results_dir / f"{self.video_id}_ocr_qwen.txt"
        self.refined_file = self.results_dir / f"{self.video_id}_refined.json"
        self.pdf_file = self.results_dir / f"{self.video_id}_frames_pro.pdf"
        self.docx_file = self.results_dir / f"{self.video_id}_frames_pro.docx"
        
        # 初始化OCR
        ocr_api_key = ocr_config.get("api_key") if ocr_config else None
        ocr_model = ocr_config.get("model", "qwen-vl-max") if ocr_config else "qwen-vl-max"
        try:
            self.ocr = QwenOCR(api_key=ocr_api_key, model=ocr_model)
        except Exception as e:
            print(f"警告: OCR初始化失败 - {e}")
            self.ocr = None
        
        # 初始化LLM Agent
        if use_llm and llm_config:
            try:
                self.llm_agent = TextRefineAgent(
                    api_key=llm_config.get("api_key"),
                    base_url=llm_config.get("base_url"),
                    model=llm_config.get("model"),
                    use_batch_mode=True  # 启用批处理模式
                )
            except Exception as e:
                print(f"警告: LLM初始化失败 - {e}")
                self.llm_agent = None
        else:
            self.llm_agent = None
        
        # 初始化ASR (Qwen)
        self.asr_config = asr_config or {}
        if asr_engine == "qwen":
            try:
                self.qwen_asr = QwenASR(
                    api_key=self.asr_config.get("api_key"),
                    region=self.asr_config.get("region", "intl")
                )
            except Exception as e:
                print(f"警告: Qwen ASR初始化失败 - {e}")
                self.qwen_asr = None
        else:
            self.qwen_asr = None
        
        # Whisper模型 (延迟加载)
        self.whisper_model = None
        
        print(f"\n{'='*80}")
        print(f"{'视频分析器 Pro 版':^80}")
        print(f"{'='*80}")
        print(f"📹 视频文件: {video_file}")
        print(f"📊 视频ID: {self.video_id}")
        print(f"💾 输出目录: {self.results_dir}")
        print(f"🎤 语音引擎: {asr_engine.upper()}")
        print(f"🤖 使用LLM: {'是' if use_llm else '否'}")
        print(f"⏭️  跳过已完成: {'是' if skip_existing else '否'}")
        print(f"{'='*80}\n")
    
    def run_full_pipeline(self):
        """运行完整处理流程"""
        try:
            # 步骤1: 提取关键帧
            self._step1_extract_frames()
            
            # 步骤2: 提取音频
            self._step2_extract_audio()
            
            # 步骤3: 语音转文字
            self._step3_transcribe_audio()
            
            # 步骤4: OCR识别
            self._step4_ocr_recognition()
            
            # 步骤5: LLM精炼
            if self.use_llm:
                self._step5_llm_refine()
            else:
                print(f"\n{'─'*80}")
                print(f"【步骤 5/6】LLM精炼整合文字")
                print(f"{'─'*80}")
                print("⏭️  跳过LLM精炼 (--no-llm)")
                # 创建简单的refined文件
                self._create_simple_refined()
            
            # 步骤6: 生成PDF和Word
            self._step6_generate_pdf()
            
            # 完成总结
            self._print_summary()
            
        except KeyboardInterrupt:
            print("\n\n⚠️  用户中断处理")
            print("💡 提示: 再次运行可从断点继续处理")
            sys.exit(0)
        except Exception as e:
            print(f"\n\n❌ 处理失败: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    
    def _step1_extract_frames(self):
        """步骤1: 提取视频关键帧"""
        print(f"\n{'─'*80}")
        print(f"【步骤 1/6】提取视频关键帧")
        print(f"{'─'*80}")
        
        # 检查是否已存在
        existing_frames = list(self.frames_dir.glob("frame_*.jpg"))
        if self.skip_existing and existing_frames:
            print(f"✓ 发现已提取的帧: {len(existing_frames)} 张 (跳过)")
            return
        
        # 清空帧目录
        for f in existing_frames:
            f.unlink()
        
        # 提取关键帧
        print(f"⏳ 正在提取关键帧...")
        frame_count = self._extract_key_frames()
        print(f"✓ 成功提取 {frame_count} 个关键帧")
    
    def _extract_key_frames(self, 
                           scene_threshold: int = 30,
                           min_interval: float = 5.0) -> int:
        """提取视频关键帧"""
        video = cv2.VideoCapture(self.video_file)
        fps = video.get(cv2.CAP_PROP_FPS)
        total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = total_frames / fps
        
        print(f"📊 视频时长: {duration:.1f}秒 ({int(duration//60)}分{int(duration%60)}秒)")
        print(f"🎬 总帧数: {total_frames}")
        print(f"📹 帧率: {fps:.2f} FPS")
        
        prev_frame = None
        prev_time = -min_interval
        frame_idx = 0
        saved_count = 0
        
        while True:
            ret, frame = video.read()
            if not ret:
                break
            
            current_time = frame_idx / fps
            
            # 转为灰度图用于比较
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            
            # 检查是否需要保存
            should_save = False
            
            if prev_frame is None:
                # 第一帧
                should_save = True
            elif current_time - prev_time >= min_interval:
                # 计算差异
                diff = cv2.absdiff(gray, prev_frame)
                diff_score = np.mean(diff)
                
                if diff_score > scene_threshold:
                    should_save = True
            
            if should_save:
                # 保存帧
                filename = f"frame_{saved_count:04d}_{current_time:.2f}s.jpg"
                filepath = self.frames_dir / filename
                cv2.imwrite(str(filepath), frame)
                
                prev_frame = gray
                prev_time = current_time
                saved_count += 1
            
            frame_idx += 1
        
        video.release()
        return saved_count
    
    def _step2_extract_audio(self):
        """步骤2: 提取视频音频"""
        print(f"\n{'─'*80}")
        print(f"【步骤 2/6】提取视频音频")
        print(f"{'─'*80}")
        
        if self.skip_existing and self.audio_file.exists():
            file_size = self.audio_file.stat().st_size / 1024 / 1024
            print(f"✓ 发现已提取的音频: {self.audio_file.name} ({file_size:.1f} MB) (跳过)")
            return
        
        print(f"⏳ 正在从视频提取音频...")
        video = VideoFileClip(self.video_file)
        video.audio.write_audiofile(str(self.audio_file), logger=None)
        
        file_size = self.audio_file.stat().st_size / 1024 / 1024
        print(f"✓ 音频提取成功: {self.audio_file}")
        print(f"  文件大小: {file_size:.1f} MB")
    
    def _step3_transcribe_audio(self):
        """步骤3: 语音转文字识别"""
        print(f"\n{'─'*80}")
        print(f"【步骤 3/6】语音转文字识别")
        print(f"{'─'*80}")
        
        if self.skip_existing and self.transcript_file.exists():
            print(f"✓ 发现已有转录文件: {self.transcript_file.name} (跳过)")
            return
        
        if self.asr_engine == "qwen":
            # 使用Qwen ASR
            self._transcribe_with_qwen()
        else:
            # 使用Whisper
            self._transcribe_with_whisper()
    
    def _transcribe_with_whisper(self):
        """使用Whisper进行转录"""
        # 加载Whisper模型
        if self.whisper_model is None:
            print(f"⏳ 正在加载 Whisper 模型 ({self.whisper_model_name})...")
            self.whisper_model = whisper.load_model(self.whisper_model_name)
            print(f"✓ 模型加载成功")
        
        # 转录
        print(f"⏳ 正在转录音频 (这可能需要几分钟)...")
        result = self.whisper_model.transcribe(str(self.audio_file), language="zh")
        
        # 保存转录结果
        with open(self.transcript_file, 'w', encoding='utf-8') as f:
            for segment in result['segments']:
                start = segment['start']
                end = segment['end']
                text = segment['text'].strip()
                f.write(f"[{start:.2f}s - {end:.2f}s] {text}\n")
        
        print(f"✓ 转录完成: {self.transcript_file}")
        print(f"  总段落: {len(result['segments'])} 个")
    
    def _transcribe_with_qwen(self):
        """使用Qwen ASR进行转录"""
        if not self.qwen_asr:
            print("❌ Qwen ASR未初始化")
            return
        
        print(f"⏳ 正在使用 Qwen ASR 转录音频...")
        
        language = self.asr_config.get("language", "zh")
        enable_itn = self.asr_config.get("enable_itn", False)
        
        # 转录
        text = self.qwen_asr.transcribe_audio(
            str(self.audio_file),
            language=language,
            enable_itn=enable_itn
        )
        
        if text:
            # 保存转录结果 (Qwen ASR返回完整文本,没有时间戳)
            with open(self.transcript_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"✓ 转录完成: {self.transcript_file}")
            print(f"  总字符数: {len(text)}")
        else:
            print(f"❌ 转录失败")
    
    def _step4_ocr_recognition(self):
        """步骤4: OCR识别课件文字"""
        print(f"\n{'─'*80}")
        print(f"【步骤 4/6】OCR识别课件文字")
        print(f"{'─'*80}")
        
        if self.skip_existing and self.ocr_file.exists():
            print(f"✓ 发现已有OCR文件: {self.ocr_file.name} (跳过)")
            return
        
        if not self.ocr:
            print("❌ OCR未初始化")
            return
        
        # 获取所有帧
        frames = sorted(self.frames_dir.glob("frame_*.jpg"))
        print(f"⏳ 正在识别 {len(frames)} 张图片...")
        print(f"   使用模型: 通义千问OCR")
        
        # 批量识别
        results = []
        success_count = 0
        
        for idx, frame_path in enumerate(frames, 1):
            print(f"\r识别进度: {idx}/{len(frames)}", end="", flush=True)
            
            text = self.ocr.recognize_image(str(frame_path))
            if text and text.strip():
                results.append(f"\n{'='*60}")
                results.append(f"文件: {frame_path.name}")
                results.append(f"{'='*60}")
                results.append(text)
                success_count += 1
            else:
                results.append(f"\n{'='*60}")
                results.append(f"文件: {frame_path.name}")
                results.append(f"{'='*60}")
                results.append("(无法识别)")
        
        print()  # 换行
        
        # 保存结果
        with open(self.ocr_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(results))
        
        print(f"✓ OCR识别完成: {self.ocr_file}")
        print(f"  成功识别: {success_count}/{len(frames)} 张")
    
    def _step5_llm_refine(self):
        """步骤5: LLM精炼整合文字"""
        print(f"\n{'─'*80}")
        print(f"【步骤 5/6】LLM精炼整合文字")
        print(f"{'─'*80}")
        
        if self.skip_existing and self.refined_file.exists():
            print(f"✓ 发现已有精炼文件: {self.refined_file.name} (跳过)")
            return
        
        if not self.llm_agent:
            print("⚠️  LLM未配置,跳过精炼")
            self._create_simple_refined()
            return
        
        # 读取转录文本
        transcript_segments = self._load_transcript()
        
        # 读取OCR结果
        ocr_texts = self._load_ocr_results()
        
        # 获取所有帧及其时间
        frames = sorted(self.frames_dir.glob("frame_*.jpg"))
        image_times = []
        for frame_path in frames:
            time_str = frame_path.name.split('_')[-1].replace('s.jpg', '')
            frame_time = float(time_str)
            image_times.append((frame_time, str(frame_path)))
        
        print(f"⏳ 正在精炼 {len(frames)} 个片段...")
        
        # 使用批处理模式精炼
        if self.llm_agent and self.llm_agent.client:
            try:
                refined_results = self.llm_agent.batch_refine_all_frames(
                    transcript_segments,
                    ocr_texts,
                    image_times
                )
            except Exception as e:
                print(f"⚠️  批处理失败: {e}, 使用简单合并")
                refined_results = self._simple_merge_results(transcript_segments, ocr_texts, frames)
        else:
            refined_results = self._simple_merge_results(transcript_segments, ocr_texts, frames)
        
        # 转换为frame_name为key的格式
        refined_by_name = {}
        for frame_path, refined_text in refined_results.items():
            if isinstance(frame_path, str):
                frame_name = Path(frame_path).name
            else:
                frame_name = frame_path.name
            refined_by_name[frame_name] = refined_text
        
        # 保存精炼结果
        with open(self.refined_file, 'w', encoding='utf-8') as f:
            json.dump(refined_by_name, f, ensure_ascii=False, indent=2)
        
        print(f"✓ 精炼完成: {self.refined_file}")
    
    def _simple_merge_results(self, transcript_segments, ocr_texts, frames):
        """简单合并转录和OCR结果(不使用LLM)"""
        results = {}
        for frame_path in frames:
            frame_name = frame_path.name
            time_str = frame_name.split('_')[-1].replace('s.jpg', '')
            frame_time = float(time_str)
            
            # 获取对应时间段的转录
            frame_segments = self._get_segments_for_time(transcript_segments, frame_time)
            transcript_text = " ".join([text for _, _, text in frame_segments])
            
            # 获取OCR文本
            ocr_text = ocr_texts.get(frame_name, "")
            
            # 简单合并
            results[frame_name] = transcript_text or ocr_text
        
        return results
    
    def _create_simple_refined(self):
        """创建简单的refined文件(不使用LLM)"""
        # 读取转录文本
        transcript_segments = self._load_transcript()
        
        # 读取OCR结果
        ocr_texts = self._load_ocr_results()
        
        # 获取所有帧
        frames = sorted(self.frames_dir.glob("frame_*.jpg"))
        
        # 简单合并
        refined_results = {}
        
        for frame_path in frames:
            frame_name = frame_path.name
            time_str = frame_name.split('_')[-1].replace('s.jpg', '')
            frame_time = float(time_str)
            
            ocr_text = ocr_texts.get(frame_name, "")
            frame_segments = self._get_segments_for_time(transcript_segments, frame_time)
            transcript_text = " ".join([text for _, _, text in frame_segments])
            
            # 优先使用转录,其次OCR
            refined_text = transcript_text if transcript_text.strip() else ocr_text
            refined_results[frame_name] = refined_text
        
        # 保存
        with open(self.refined_file, 'w', encoding='utf-8') as f:
            json.dump(refined_results, f, ensure_ascii=False, indent=2)
        
        print(f"✓ 创建简单合并结果: {self.refined_file}")
    
    def _step6_generate_pdf(self):
        """步骤6: 生成PDF和Word文件"""
        print(f"\n{'─'*80}")
        print(f"【步骤 6/6】生成PDF和Word文件")
        print(f"{'─'*80}")
        
        # 读取精炼结果
        if not self.refined_file.exists():
            print("❌ 未找到精炼结果文件")
            return
        
        with open(self.refined_file, 'r', encoding='utf-8') as f:
            refined_data = json.load(f)
        
        # 获取所有帧
        frames = sorted(self.frames_dir.glob("frame_*.jpg"))
        
        print(f"⏳ 正在生成PDF...")
        print(f"   图片数量: {len(frames)}")
        print(f"   使用文字: {'LLM精炼' if self.use_llm else '简单合并'}")
        
        # 生成PDF
        self._generate_pdf_document(frames, refined_data)
        
        # 生成Word文档
        if HAS_DOCX:
            self._generate_word_document(frames, refined_data)
    
    def _generate_pdf_document(self, frames: List[Path], refined_data: Dict):
        """生成PDF文档"""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.utils import ImageReader
        
        # 注册中文字体
        try:
            pdfmetrics.registerFont(TTFont('SimSun', '/System/Library/Fonts/STHeiti Light.ttc'))
            font_name = 'SimSun'
        except:
            try:
                pdfmetrics.registerFont(TTFont('SimSun', '/System/Library/Fonts/PingFang.ttc'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'
        
        # 创建PDF
        c = canvas.Canvas(str(self.pdf_file), pagesize=A4)
        page_width, page_height = A4
        
        for frame_path in frames:
            frame_name = frame_path.name
            text = refined_data.get(frame_name, "")
            
            # 添加图片页
            img = Image.open(frame_path)
            img_width, img_height = img.size
            
            # 计算缩放比例
            scale = min((page_width - 40) / img_width, (page_height - 100) / img_height)
            new_width = img_width * scale
            new_height = img_height * scale
            
            x = (page_width - new_width) / 2
            y = page_height - new_height - 50
            
            c.drawImage(ImageReader(img), x, y, new_width, new_height)
            
            # 添加标题
            c.setFont(font_name, 10)
            c.drawString(50, page_height - 30, frame_name)
            
            c.showPage()
            
            # 添加文字页
            if text and text.strip():
                c.setFont(font_name, 12)
                
                # 标题
                c.drawString(50, page_height - 50, f"文字内容: {frame_name}")
                
                # 分段显示文字
                y_position = page_height - 100
                max_width = page_width - 100
                
                lines = self._wrap_text(text, max_width, font_name, 12)
                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = page_height - 50
                    
                    c.drawString(50, y_position, line)
                    y_position -= 20
                
                c.showPage()
        
        c.save()
        
        file_size = self.pdf_file.stat().st_size / 1024 / 1024
        page_count = len(frames) * 2
        print(f"✓ PDF生成成功: {self.pdf_file}")
        print(f"  总页数: {page_count} 页")
        print(f"  文件大小: {file_size:.1f} MB")
    
    def _generate_word_document(self, frames: List[Path], refined_data: Dict):
        """生成Word文档"""
        print(f"⏳ 正在生成Word文档...")
        
        doc = Document()
        
        # 添加标题
        doc.add_heading(f'视频分析报告 - {self.video_id}', 0)
        
        # 添加每一帧
        for frame_path in frames:
            frame_name = frame_path.name
            text = refined_data.get(frame_name, "")
            
            # 添加帧标题
            doc.add_heading(frame_name, level=2)
            
            # 添加图片
            try:
                doc.add_picture(str(frame_path), width=Inches(6))
            except Exception as e:
                print(f"  警告: 无法添加图片 {frame_name}: {e}")
            
            # 添加文字内容
            if text and text.strip():
                # 处理Markdown格式
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    else:
                        doc.add_paragraph(line)
            
            # 添加分页符
            doc.add_page_break()
        
        # 保存
        doc.save(str(self.docx_file))
        
        file_size = self.docx_file.stat().st_size / 1024 / 1024
        print(f"✓ Word文档生成成功: {self.docx_file}")
        print(f"  文件大小: {file_size:.1f} MB")
    
    def _wrap_text(self, text: str, max_width: float, font_name: str, font_size: int) -> List[str]:
        """文本换行"""
        from reportlab.pdfbase import pdfmetrics
        
        lines = []
        current_line = ""
        
        for char in text:
            test_line = current_line + char
            width = pdfmetrics.stringWidth(test_line, font_name, font_size)
            
            if width > max_width:
                if current_line:
                    lines.append(current_line)
                current_line = char
            else:
                current_line = test_line
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _load_transcript(self) -> List[Tuple[float, float, str]]:
        """加载转录文本"""
        if not self.transcript_file.exists():
            return []
        
        segments = []
        
        with open(self.transcript_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                
                # 解析时间戳 [start - end] text
                if line.startswith('['):
                    try:
                        time_part = line[1:line.index(']')]
                        text_part = line[line.index(']')+1:].strip()
                        
                        if ' - ' in time_part:
                            start_str, end_str = time_part.split(' - ')
                            start = float(start_str.replace('s', ''))
                            end = float(end_str.replace('s', ''))
                            segments.append((start, end, text_part))
                    except:
                        pass
                else:
                    # Qwen ASR格式 (无时间戳)
                    segments.append((0.0, 0.0, line))
        
        return segments
    
    def _load_ocr_results(self) -> Dict[str, str]:
        """加载OCR结果 - 支持旧格式和Qwen OCR格式"""
        if not self.ocr_file.exists():
            return {}
        
        results = {}
        current_file = None
        current_text = []
        
        with open(self.ocr_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                
                # 旧格式: "文件: xxx"
                if line.startswith('文件: '):
                    # 保存上一个文件的结果
                    if current_file and current_text:
                        results[current_file] = '\n'.join(current_text)
                    
                    # 开始新文件
                    current_file = line.replace('文件: ', '')
                    current_text = []
                
                # 新格式(Qwen OCR): "【frame_xxx】"
                elif line.startswith('【') and line.endswith('】'):
                    # 保存上一个文件的结果
                    if current_file and current_text:
                        results[current_file] = '\n'.join(current_text)
                    
                    # 开始新文件
                    current_file = line[1:-1]  # 去掉【】
                    current_text = []
                
                # 跳过分隔线和空行
                elif line and not line.startswith('=') and not line.startswith('-') and line != '(无法识别)' and line != '通义千问 OCR 识别结果':
                    current_text.append(line)
        
        # 保存最后一个文件
        if current_file and current_text:
            results[current_file] = '\n'.join(current_text)
        
        return results
    
    def _get_segments_for_time(self, 
                               segments: List[Tuple[float, float, str]], 
                               frame_time: float,
                               window: float = 30.0) -> List[Tuple[float, float, str]]:
        """获取指定时间附近的转录段落"""
        result = []
        
        for start, end, text in segments:
            # Qwen ASR格式 (无时间戳)
            if start == 0.0 and end == 0.0:
                result.append((start, end, text))
            # Whisper格式
            elif abs(start - frame_time) <= window or (start <= frame_time <= end):
                result.append((start, end, text))
        
        return result
    
    def _print_summary(self):
        """打印处理完成总结"""
        frames = list(self.frames_dir.glob("frame_*.jpg"))
        
        print(f"\n{'='*80}")
        print(f"{'处理完成总结':^80}")
        print(f"{'='*80}")
        print(f"📹 视频文件: {self.video_file}")
        print(f"🖼️  提取帧数: {len(frames)} 个")
        
        if self.pdf_file.exists():
            pdf_size = self.pdf_file.stat().st_size / 1024 / 1024
            print(f"📄 PDF文件: {self.pdf_file} ({pdf_size:.1f} MB)")
        
        if self.docx_file.exists():
            docx_size = self.docx_file.stat().st_size / 1024 / 1024
            print(f"📝 Word文件: {self.docx_file} ({docx_size:.1f} MB)")
        
        print(f"💾 所有结果: {self.results_dir}/")
        print(f"{'='*80}")
        print(f"✅ 所有步骤已完成!")
        print(f"{'='*80}\n")


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="视频分析器 Pro 版 - 完整自动化处理",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基础用法
  python3 -m video_analyzer.core.analyzer --video lecture.mp4
  
  # 使用Qwen ASR
  python3 -m video_analyzer.core.analyzer --video lecture.mp4 --asr-engine qwen
  
  # 不使用LLM精炼
  python3 -m video_analyzer.core.analyzer --video lecture.mp4 --no-llm
  
  # 强制重新处理
  python3 -m video_analyzer.core.analyzer --video lecture.mp4 --no-skip
        """
    )
    
    parser.add_argument('--video', required=True, help='视频文件路径')
    parser.add_argument('--whisper-model', default='base', 
                       choices=['tiny', 'base', 'small', 'medium', 'large'],
                       help='Whisper模型大小 (默认: base)')
    parser.add_argument('--no-llm', action='store_true', help='不使用LLM精炼')
    parser.add_argument('--no-skip', action='store_true', help='不跳过已存在的文件,重新处理')
    
    # ASR配置
    parser.add_argument('--asr-engine', default='whisper',
                       choices=['whisper', 'qwen'],
                       help='语音识别引擎 (默认: whisper)')
    parser.add_argument('--asr-api-key', help='ASR API Key (Qwen ASR使用)')
    parser.add_argument('--asr-region', default='intl',
                       choices=['intl', 'cn'],
                       help='Qwen ASR区域 (默认: intl)')
    parser.add_argument('--asr-language', default='zh', help='语音识别语言 (默认: zh)')
    parser.add_argument('--asr-enable-itn', action='store_true', 
                       help='启用ITN(逆文本正则化)')
    
    # LLM配置
    parser.add_argument('--api-key', help='LLM API Key')
    parser.add_argument('--base-url', help='LLM API Base URL')
    parser.add_argument('--model',  help='LLM模型名称')
    
    # OCR配置
    parser.add_argument('--ocr-api-key', help='通义千问API Key')
    parser.add_argument('--ocr-model', default='qwen-vl-max', 
                       choices=['qwen-vl-max', 'qwen-vl-plus', 'qwen3-vl-plus'],
                       help='OCR模型名称')
    
    args = parser.parse_args()
    
    # LLM配置
    llm_config = {
        "api_key": args.api_key,
        "base_url": args.base_url,
        "model": args.model
    }
    
    # OCR配置
    ocr_config = {
        "api_key": args.ocr_api_key,
        "model": args.ocr_model
    }
    
    # ASR配置
    asr_config = {
        "api_key": args.asr_api_key,
        "region": args.asr_region,
        "language": args.asr_language,
        "enable_itn": args.asr_enable_itn
    }
    
    # 创建分析器
    try:
        analyzer = VideoAnalyzerPro(
            video_file=args.video,
            use_llm=not args.no_llm,
            llm_config=llm_config,
            ocr_config=ocr_config,
            whisper_model=args.whisper_model,
            asr_engine=args.asr_engine,
            asr_config=asr_config,
            skip_existing=not args.no_skip
        )
        
        # 运行完整流程
        analyzer.run_full_pipeline()
        
    except FileNotFoundError as e:
        print(f"\n❌ 错误: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ 初始化失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
