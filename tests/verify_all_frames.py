#!/usr/bin/env python3
"""验证所有帧的结构"""

from docx import Document

doc = Document('results/1732328440969754_frames_pro.docx')

print('验证所有帧的结构')
print('=' * 70)

frame_count = 0
correct_order_count = 0
errors = []

i = 0
while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    # 查找帧标题
    if text.startswith('帧 '):
        frame_count += 1
        frame_title = text
        
        # 检查下一段是否是图片
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            has_image = any(run._element.xpath('.//pic:pic') for run in next_para.runs)
            
            if has_image:
                # 检查图片后是否是文字标题
                if i + 2 < len(doc.paragraphs):
                    text_header = doc.paragraphs[i + 2].text.strip()
                    if '文字内容' in text_header or '[此帧无文字内容]' in text_header:
                        correct_order_count += 1
                        if frame_count <= 3:
                            print(f'✓ {frame_title}: 图片 → 文字 (正确)')
                    else:
                        errors.append(f'{frame_title}: 图片后不是文字标题')
            else:
                errors.append(f'{frame_title}: 帧标题后不是图片')
    
    i += 1

print()
print(f'总帧数: {frame_count}')
print(f'顺序正确: {correct_order_count}')
print(f'顺序错误: {len(errors)}')

if errors:
    print('\n错误列表:')
    for err in errors[:5]:
        print(f'  - {err}')
else:
    print('\n✅ 所有帧的顺序都正确!')

# 随机抽查几帧
print('\n随机抽查 (第10, 20, 30帧):')
print('-' * 70)

check_frames = [10, 20, 30]
for frame_num in check_frames:
    target = f'帧 {frame_num}/'
    for i, para in enumerate(doc.paragraphs):
        if target in para.text:
            # 检查后续3段
            print(f'\n帧 {frame_num}:')
            print(f'  {i}. [标题] {para.text.strip()}')
            if i + 1 < len(doc.paragraphs):
                has_img = any(run._element.xpath('.//pic:pic') for run in doc.paragraphs[i + 1].runs)
                print(f'  {i+1}. [图片] {"✓ 包含图片" if has_img else "✗ 无图片"}')
            if i + 2 < len(doc.paragraphs):
                print(f'  {i+2}. [文字] {doc.paragraphs[i + 2].text.strip()[:50]}...')
            break
