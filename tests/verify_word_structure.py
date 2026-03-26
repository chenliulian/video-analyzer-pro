#!/usr/bin/env python3
"""验证Word文档结构"""

from docx import Document

doc = Document('results/1732328440969754_frames_pro.docx')

print('Word文档结构验证')
print('=' * 70)

# 统计元素
total_paragraphs = len(doc.paragraphs)
total_images = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run._element.xpath('.//pic:pic'):
            total_images += 1

print(f'总段落数: {total_paragraphs}')
print(f'总图片数: {total_images}')
print()

# 查看第一帧的结构
print('第一帧结构分析:')
print('-' * 70)

frame_found = False
image_found = False
text_found = False

for i, para in enumerate(doc.paragraphs[:25]):
    text = para.text.strip()
    has_image = any(run._element.xpath('.//pic:pic') for run in para.runs)
    
    if '帧 1/36' in text:
        print(f'{i}. [帧标题] {text}')
        frame_found = True
    elif has_image and frame_found and not image_found:
        print(f'{i}. [图片] (包含图片)')
        image_found = True
    elif '文字内容' in text and image_found and not text_found:
        print(f'{i}. [文字标题] {text}')
        text_found = True
    elif text_found and text and len(text) > 10:
        preview = text[:100] + '...' if len(text) > 100 else text
        print(f'{i}. [文字内容] {preview}')
        break

print()
if frame_found and image_found and text_found:
    print('✅ 顺序验证: 正确 (帧标题 → 图片 → 文字内容)')
else:
    print('⚠️  顺序验证: 需要检查')
    print(f'   帧标题: {frame_found}, 图片: {image_found}, 文字: {text_found}')
